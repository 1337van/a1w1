import streamlit as st
import os
import tempfile
import base64
import re
import subprocess
import imageio_ffmpeg as iio_ffmpeg
from google import genai
from google.genai.types import HttpOptions, Part
from google.cloud import storage
from docx import Document
from docx.shared import Inches

# --- CONFIG & CLIENT SETUP ---
cfg = st.secrets["gcp"]
PROJECT_ID, LOCATION, BUCKET, SA_BASE64 = (
    cfg["project"], cfg["location"], cfg["bucket"], cfg["sa_key"]
)
tmp_dir = tempfile.mkdtemp()
sa_path = os.path.join(tmp_dir, "sa.json")
with open(sa_path, "wb") as f:
    f.write(base64.b64decode(SA_BASE64))
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = sa_path
os.environ["GOOGLE_CLOUD_PROJECT"]       = PROJECT_ID
os.environ["GOOGLE_CLOUD_LOCATION"]      = LOCATION
os.environ["GOOGLE_GENAI_USE_VERTEXAI"]  = "True"

client = genai.Client(http_options=HttpOptions(api_version="v1"))
storage_client = storage.Client()
FFMPEG = iio_ffmpeg.get_ffmpeg_exe()

# --- PAGE LAYOUT & CSS ---
st.set_page_config(page_title="📦 Video-to-WI Generator", layout="centered")
st.markdown("""
<style>
  .main .block-container {
    max-width:700px; margin:auto; padding:2rem 1rem; background:#f0f4f8;
  }
  button[kind="primary"] {
    background-color:#0057a6!important; border-color:#0057a6!important;
  }
</style>
""", unsafe_allow_html=True)

# --- HEADER ---
st.image("https://i.postimg.cc/L8JXmQ7t/gwlogo1.jpg", width=120)
st.title("📦 Video Summarizer → Work Instructions")
st.caption("powered by Vertex AI Flash 2.0")
st.markdown("---")

# --- FULL PROMPT 1.0–7.0 ---
PROMPT = """\
You are an operations specialist with a background in quality analysis and engineering technician practices, observing a manufacturing process within a controlled ISO 9001:2015 environment.

Visually and audibly analyze the video input to generate structured work instructions.

**For each action step, prefix the line with the video timestamp formatted exactly as `[MM:SS]`.**

Follow this output template format:

1.0 Purpose  
Describe the purpose of this work instruction.

2.0 Scope  
State the scope of this procedure (e.g., "This applies to Goodwill Commercial Services").

3.0 Responsibilities  
ROLE     | RESPONSIBILITY  
-------- | ----------------  
Line Lead | ● Ensure procedural adherence, documentation, and nonconformance decisions  
Operator  | ● Follow instructions and execute the defined procedure

4.0 Tools, Materials, Equipment, Supplies  
DESCRIPTION | VISUAL | HAZARD  
----------- | ------ | ------  
(e.g. Box Cutter | [insert image] | Sharp Blade Hazard)

5.0 Associated Safety and Ergonomic Concerns  
List relevant safety issues; include legend if needed.

6.0 Procedure  
STEP | TIMESTAMP | ACTION | VISUAL | HAZARD  
-----|-----------|--------|--------|-------  
1    | `[MM:SS]` | Describe action clearly | [frame] | [Identify hazard]  
2    | `[MM:SS]` | Continue for each step    | [frame] | [Identify hazard]

> If any part of the process is unclear, mark as **[uncertain action]**.

7.0 Reference Documents  
List any applicable SOPs, work orders, or specs.

Keep formatting clean and ensure each action step ties to its visual frame.
"""
prompt = st.text_area("Prompt (1.0–7.0):", PROMPT, height=320)
st.markdown("---")

# --- UPLOAD & GENERATE ---
video = st.file_uploader("Upload .mp4 video", type="mp4")
if video and st.button("Generate Draft WI", type="primary"):
    # Save locally
    local = os.path.join(tmp_dir, video.name)
    with open(local, "wb") as f:
        f.write(video.read())

    # Upload to GCS
    gcs = f"input/{video.name}"
    storage_client.bucket(BUCKET).blob(gcs).upload_from_filename(local)
    st.success("Uploaded; generating…")

    # Call Vertex AI
    with st.spinner("Calling Vertex AI…"):
        resp = client.models.generate_content(
            model="gemini-2.0-flash-001",
            contents=[
                Part.from_uri(file_uri=f"gs://{BUCKET}/{gcs}", mime_type="video/mp4"),
                prompt
            ],
        )

    # Raw response
    raw = resp.text
    st.session_state.raw = raw

    # Strip any leading boilerplate before "1.0 Purpose"
    lines = raw.splitlines()
    for idx, line in enumerate(lines):
        if re.match(r"^1\.0\s+Purpose", line):
            cleaned = "\n".join(lines[idx:])
            break
    else:
        cleaned = raw
    st.session_state.summary = cleaned

    # Display cleaned draft
    st.markdown("#### Full Draft (Sections 1.0–7.0)")
    st.code(cleaned, language="markdown")

    # Parse 6.0 Procedure rows
    steps = []
    in_proc = False
    for line in cleaned.splitlines():
        if re.match(r"^6\.0\s+Procedure", line):
            in_proc = True
            continue
        if in_proc and line.strip().startswith("| ["):
            cols = [c.strip() for c in line.split("|")[1:-1]]
            # cols = [STEP, TIMESTAMP, ACTION, VISUAL, HAZARD]
            ts = re.search(r"\[(\d{2}:\d{2})\]", cols[1]).group(1)
            action = cols[2]
            hazard = cols[4]
            steps.append((ts, action, hazard))
        if in_proc and line.strip() == "":
            break
    st.session_state.steps = steps

    # Extract frames for each timestamp
    frames = []
    for ts, _, _ in steps:
        img = os.path.join(tmp_dir, f"frame_{ts.replace(':','_')}.png")
        subprocess.run(
            [FFMPEG, "-y", "-ss", ts, "-i", local, "-vframes", "1", img],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        if os.path.exists(img):
            frames.append({"time": ts, "path": img})
    st.session_state.frames = frames

st.markdown("---")

# --- IMAGE REVIEW via selectbox ---
if "summary" in st.session_state:
    if st.session_state.frames:
        labels = [
            f"Step {i+1} – [{ts}] {act}"
            for i, (ts, act, _) in enumerate(st.session_state.steps)
        ]
        choice = st.selectbox("Select a step to preview", labels)
        idx = labels.index(choice)
        ts, action, hazard = st.session_state.steps[idx]
        img = st.session_state.frames[idx]["path"]

        st.markdown(f"**{choice}**  \n_Hazard:_ {hazard}")
        st.image(img, use_container_width=True)
        if st.button("Delete this image"):
            st.session_state.steps.pop(idx)
            st.session_state.frames.pop(idx)
            st.experimental_rerun()
    else:
        st.info("No Procedure steps found or no frames extracted.")

    st.markdown("---")

    # --- DOCX EXPORT (1–7 + images in 6.0) ---
    if st.button("Download WI .docx"):
        doc = Document()
        # Write sections 1.0–5.0 and 7.0
        lines = st.session_state.summary.splitlines()
        i = 0
        while i < len(lines):
            m = re.match(r"^(\d\.\d)\s+(.*)", lines[i])
            if m and m.group(1) != "6.0":
                # Section header
                doc.add_heading(lines[i], level=1)
                i += 1
                # Section body
                while i < len(lines) and not re.match(r"^\d\.\d", lines[i]):
                    if lines[i].strip():
                        doc.add_paragraph(lines[i])
                    i += 1
                continue
            if m and m.group(1) == "6.0":
                doc.add_heading("6.0 Procedure", level=1)
                # Skip to the first table row
                while i < len(lines) and not lines[i].strip().startswith("|"):
                    i += 1
                # Skip header + separator
                i += 2
                # Insert parsed steps
                for idx, (ts, act, haz) in enumerate(st.session_state.steps, start=1):
                    p = doc.add_paragraph(style="Heading 2")
                    p.add_run(f"Step {idx} – [{ts}] {act}")
                    doc.add_paragraph(f"Hazard: {haz}", style="List Bullet")
                    img = next((f["path"] for f in st.session_state.frames if f["time"] == ts), None)
                    if img:
                        doc.add_picture(img, width=Inches(3))
                        doc.add_paragraph()
                break
            i += 1

        # Ensure 7.0 Reference Documents exists
        if not any(l.startswith("7.0") for l in lines):
            doc.add_heading("7.0 Reference Documents", level=1)

        out_path = os.path.join(tmp_dir, "Work_Instructions.docx")
        doc.save(out_path)
        with open(out_path, "rb") as f:
            st.download_button("Download WI .docx", f, file_name="Work_Instructions.docx")
